"""
DOCX generation service for Resume-Matcher.

Converts optimized resume text into professionally formatted .docx files.
"""

import io
import logging

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)


class DOCXGenerationError(Exception):
    """Raised when DOCX generation fails."""

    pass


class DOCXGenerationService:
    """Service for generating professionally formatted DOCX resumes."""

    def __init__(self):
        """Initialize DOCX generation service."""
        # Default formatting settings
        self.font_name = "Calibri"
        self.font_size_normal = 11
        self.font_size_heading = 14
        self.font_size_name = 18
        self.line_spacing = 1.15

    async def generate_docx(self, optimized_text: str, user_name: str | None = None) -> bytes:
        """
        Generate a professionally formatted DOCX file from optimized resume text.

        Args:
            optimized_text: Optimized resume text from AI
            user_name: Optional user name for document properties

        Returns:
            DOCX file content as bytes

        Raises:
            DOCXGenerationError: If generation fails
        """
        try:
            logger.info(f"Generating DOCX resume (user: {user_name or 'anonymous'})")

            # Create new document
            document = Document()

            # Set document properties
            if user_name:
                document.core_properties.author = user_name
                document.core_properties.title = f"Currículo - {user_name}"

            # Set margins (1 inch all around)
            sections = document.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # Parse and format the resume text
            self._format_resume_content(document, optimized_text)

            # Save to bytes
            docx_bytes = self._save_to_bytes(document)

            logger.info(f"DOCX generated successfully - Size: {len(docx_bytes)} bytes")
            return docx_bytes

        except Exception as e:
            logger.exception(f"Error generating DOCX: {str(e)}")
            raise DOCXGenerationError(f"Falha ao gerar arquivo DOCX: {str(e)}")

    def _format_resume_content(self, document: Document, text: str) -> None:
        """
        Format resume content with professional styling.

        Args:
            document: DOCX document object
            text: Resume text to format
        """
        # Split text into lines
        lines = text.split("\n")

        for i, line in enumerate(lines):
            line = line.strip()

            if not line:
                # Add spacing between sections
                if i > 0:
                    document.add_paragraph()
                continue

            # Detect section type and apply formatting
            if self._is_name_heading(line, i):
                # Name (first significant line)
                para = document.add_paragraph(line)
                self._style_name_heading(para)

            elif self._is_section_heading(line):
                # Section headings (e.g., "EXPERIÊNCIA PROFISSIONAL", "FORMAÇÃO")
                para = document.add_paragraph(line)
                self._style_section_heading(para)

            elif self._is_subsection_heading(line):
                # Subsection headings (e.g., "Engenheiro de Software Senior", "Universidade XYZ")
                para = document.add_paragraph(line)
                self._style_subsection_heading(para)

            else:
                # Normal text
                para = document.add_paragraph(line)
                self._style_normal_text(para)

    def _is_name_heading(self, line: str, index: int) -> bool:
        """Check if line is likely the candidate's name."""
        # First non-empty line that's not all caps and has reasonable length
        return index < 3 and 5 < len(line) < 100 and not line.isupper()

    def _is_section_heading(self, line: str) -> bool:
        """Check if line is a major section heading."""
        # All uppercase lines or lines ending with colon
        return (
            line.isupper()
            or line.endswith(":")
            or any(
                keyword in line.upper()
                for keyword in [
                    "EXPERIÊNCIA",
                    "FORMAÇÃO",
                    "EDUCAÇÃO",
                    "HABILIDADES",
                    "COMPETÊNCIAS",
                    "IDIOMAS",
                    "CERTIFICAÇÕES",
                    "CONTATO",
                    "RESUMO",
                    "OBJETIVO",
                ]
            )
        )

    def _is_subsection_heading(self, line: str) -> bool:
        """Check if line is a subsection heading (job title, company, etc.)."""
        # Lines that look like titles or company names
        # Heuristic: starts with capital letter, has 10-80 chars, no lowercase start
        if not line:
            return False
        return (
            line[0].isupper()
            and 10 < len(line) < 80
            and (
                " - " in line  # e.g., "Empresa XYZ - Cargo"
                or "•" in line
                or (line.count(" ") <= 6 and not line.startswith(" "))
            )
        )

    def _style_name_heading(self, paragraph: Paragraph) -> None:
        """Apply styling for name heading."""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.runs[0]
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size_name)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

    def _style_section_heading(self, paragraph: Paragraph) -> None:
        """Apply styling for major section headings."""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.runs[0]
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size_heading)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 73, 125)  # Professional blue
        # Add bottom border
        paragraph.paragraph_format.space_after = Pt(6)

    def _style_subsection_heading(self, paragraph: Paragraph) -> None:
        """Apply styling for subsection headings."""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.runs[0]
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size_normal + 1)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        paragraph.paragraph_format.space_after = Pt(3)

    def _style_normal_text(self, paragraph: Paragraph) -> None:
        """Apply styling for normal text."""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.runs[0]
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size_normal)
        run.font.color.rgb = RGBColor(0, 0, 0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = self.line_spacing

    def _save_to_bytes(self, document: Document) -> bytes:
        """
        Save document to bytes.

        Args:
            document: DOCX document object

        Returns:
            Document content as bytes
        """
        # Save to BytesIO
        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        return file_stream.read()

    async def generate_docx_from_sections(
        self,
        sections: dict[str, str | list[str]],
        user_name: str | None = None,
    ) -> bytes:
        """
        Generate DOCX from structured sections (alternative format).

        Args:
            sections: Dict with section names and content
            user_name: Optional user name

        Returns:
            DOCX file content as bytes

        Example sections:
            {
                "name": "João Silva",
                "contact": "joao@email.com | (11) 99999-9999",
                "summary": "Profissional experiente...",
                "experience": ["Cargo 1...", "Cargo 2..."],
                "education": ["Formação 1...", "Formação 2..."],
                "skills": ["Python", "JavaScript", "SQL"]
            }
        """
        try:
            document = Document()

            # Set margins
            sections_list = document.sections
            for section in sections_list:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # Add name
            if "name" in sections:
                para = document.add_paragraph(str(sections["name"]))
                self._style_name_heading(para)
                document.add_paragraph()  # Spacing

            # Add contact info
            if "contact" in sections:
                para = document.add_paragraph(str(sections["contact"]))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.runs[0]
                run.font.name = self.font_name
                run.font.size = Pt(10)
                document.add_paragraph()  # Spacing

            # Add other sections
            section_order = ["summary", "experience", "education", "skills", "certifications", "languages"]

            for section_key in section_order:
                if section_key not in sections:
                    continue

                # Add section heading
                heading_text = section_key.replace("_", " ").upper()
                heading = document.add_paragraph(heading_text)
                self._style_section_heading(heading)

                # Add section content
                content = sections[section_key]
                if isinstance(content, list):
                    for item in content:
                        para = document.add_paragraph(f"• {item}")
                        self._style_normal_text(para)
                else:
                    para = document.add_paragraph(str(content))
                    self._style_normal_text(para)

                document.add_paragraph()  # Spacing between sections

            # Save to bytes
            return self._save_to_bytes(document)

        except Exception as e:
            logger.exception(f"Error generating structured DOCX: {str(e)}")
            raise DOCXGenerationError(f"Falha ao gerar arquivo DOCX estruturado: {str(e)}")


# Singleton instance
_docx_generation_service: DOCXGenerationService | None = None


def get_docx_generation_service() -> DOCXGenerationService:
    """Get or create the singleton DOCXGenerationService instance."""
    global _docx_generation_service

    if _docx_generation_service is None:
        _docx_generation_service = DOCXGenerationService()

    return _docx_generation_service
